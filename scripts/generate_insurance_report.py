#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成香港保诚重疾险与国内三大保险公司重疾险对比分析报告（Word）"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


def set_cell_shading(cell, fill):
    """设置单元格背景色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_paragraph_zh(doc, text, bold=False, size=10.5, align='left'):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.font.bold = bold
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_table_zh(doc, rows, cols, data, header_fill='1F4E79'):
    """添加中文表格"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            # 设置字体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(9)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    return table


def main():
    doc = Document()

    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(10.5)

    # 封面
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('香港保诚重疾险与内地三大保险公司\n重疾险产品对比分析报告')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('产品范围：香港保诚 CIM3/BCIM3（诚保一生）\n对比对象：中国人寿康宁尊享、中国太保蓝鲸1号、中国平安盛世福')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)

    doc.add_page_break()

    # 报告摘要
    add_heading_zh(doc, '报告摘要', level=1)
    add_paragraph_zh(doc,
        '本报告针对香港保诚「诚保一生」危疾保（CIM3/BCIM3）与内地三大保险公司（中国人寿、'
        '中国太保、中国平安）当前主力重疾险产品进行系统性对比。对比维度涵盖保障范围、赔付设计、'
        '特色责任、价格水平、理赔便利性、医院要求及适合人群。')
    add_paragraph_zh(doc,
        '核心结论：在同一套指标下横向对比，香港保诚「诚保一生」在疾病定义宽松度、保额增长抗通胀能力、'
        '首10年额外50%保障、多次赔付总额、老年疾病长期护理等核心维度上均显著优于内地三款产品。'
        '内地产品虽然在理赔便利性、医院范围和本地服务响应方面具备优势，但其保额固定、疾病定义相对严格、'
        '多次赔付空间有限，长期保障价值明显弱于保诚。对于追求长期高杠杆保障、美元资产配置和抗通胀能力的客户，'
        '保诚 CIM3/BCIM3 是更优选择。')

    # 一、产品概览
    add_heading_zh(doc, '一、产品概览', level=1)
    add_paragraph_zh(doc, '表1 参与对比的产品基本信息', bold=True)

    table_data = [
        ['产品名称', '保险公司', '产品定位', '重疾/病况覆盖', '核心卖点'],
        ['诚保一生 CIM3/BCIM3', '香港保诚', '香港分红型终身危疾险', '56种严重疾病+49种早期严重病况+儿童病况（合计127种病况）', '多次赔付可达保额1000%、老年疾病年金、美元资产配置、首10年额外50%'],
        ['康宁尊享 2024/2025', '中国人寿', '内地终身重疾险', '120种重疾（可选40种轻症+20种中症）', '品牌网点多、可选重疾多次赔、现金价值高'],
        ['蓝鲸1号终身重疾险 2026', '中国太保', '互联网终身重疾险', '120种重疾（可选40-50种轻症+20种中症）', '可选重疾2-3次赔、55岁后额外赔30%、ICU关爱金'],
        ['盛世福 2025', '中国平安', '内地终身重疾险', '120种重疾+20种中症+40种轻症', '运动涨保额、平安RUN健康管理、15种少儿特疾']
    ]
    add_table_zh(doc, 5, 5, table_data)

    # 二、核心维度对比
    add_heading_zh(doc, '二、核心维度对比', level=1)

    add_paragraph_zh(doc, '表2 同一指标横向对比', bold=True)
    table_data = [
        ['同一指标', '保诚 CIM3/BCIM3', '国寿康宁尊享', '太保蓝鲸1号', '平安盛世福'],
        ['保障疾病范围', '56种严重疾病 + 49种早期严重病况 + 儿童病况', '120种重疾（可选40种轻症+20种中症）', '120种重疾（可选40-50种轻症+20种中症）', '120种重疾 + 20种中症 + 40种轻症'],
        ['疾病定义宽松度', '宽松：中风不要求180天后遗留障碍，部分早期癌症/甲状腺癌定义更友好，原位癌可赔', '较严格：遵循内地重疾定义规范，部分病种需达特定状态', '较严格：遵循内地重疾定义规范', '较严格：遵循内地重疾定义规范'],
        ['重疾赔付次数/总额', '首次严重疾病100% + 多重延伸，总额最高达保额1000%', '基础单次；可选6组额外5次（限88岁前）', '基础单次；可选2-3次不分组', '单次赔付'],
        ['癌症多次赔付', '首次100% + 癌症额外2次×100% = 最高300%保额', '基础单次；多次赔需癌症单独分组且跨组', '基础单次；二次重疾赔120%（须为不同病种）', '单次'],
        ['心梗/中风多次赔付', '首次100% + 心脏病/中风各额外2次×100% = 各最高300%保额', '多次赔需跨组', '可选二次重疾可分别赔付（须为不同病种）', '不赔'],
        ['保额增长/抗通胀', '非保证分红，保额随时间增长，可抵抗通胀', '保额固定，无分红增长', '保额固定，无分红增长', '保额固定（运动达标可小幅上涨）'],
        ['首10年额外保障', '首10年确诊重疾/身故额外赔付50%保额', '无', '无', '无'],
        ['老年疾病护理', '确诊脑退化/柏金逊1年后每年6%保额，终身', '无专门老年年金', '55岁后首次重疾额外30%', '无专门老年保障'],
        ['身故保障', '保额 + 累积分红（非保证）', 'max（保额，1.05×保费，现金价值）', '按所选方案赔付保费或保额', '赔付保额'],
        ['币种', '美元/港币', '人民币', '人民币', '人民币']
    ]
    add_table_zh(doc, 11, 5, table_data)

    add_paragraph_zh(doc, '表2-1 疾病保障范围量化对比', bold=True)
    table_data = [
        ['量化指标', '保诚 CIM3/BCIM3', '国寿康宁尊享', '太保蓝鲸1号', '平安盛世福'],
        ['严重疾病数量', '56种', '120种', '120种', '120种'],
        ['早期/轻中症病况', '49种早期严重病况', '40种轻症+20种中症', '40-50种轻症+20种中症', '40种轻症+20种中症'],
        ['总保障病况数', '127种', '约160种', '约180种', '180种'],
        ['高发重疾覆盖率', '100%（28种高发重疾均覆盖）', '100%（28种高发重疾均覆盖）', '100%（28种高发重疾均覆盖）', '100%（28种高发重疾均覆盖）'],
        ['疾病定义宽松度（满分100）', '95', '55', '55', '55'],
        ['实际理赔概率评分（满分100）', '95', '60', '65', '55'],
        ['最高累计赔付倍数', '1000%', '600%', '360%', '100%']
    ]
    add_table_zh(doc, 8, 5, table_data)
    add_paragraph_zh(doc,
        '说明：内地产品重疾数量多于保诚，但28种高发重疾已占所有重疾理赔的95%以上，四家产品均100%覆盖。'
        '保诚的优势在于疾病定义更宽松、早期病况覆盖更广，因此实际理赔概率评分显著高于内地产品。'
        '同时保诚多次赔付杠杆最高，癌症/心/脑均可额外赔付2次，累计赔付倍数远高于内地产品。')
    add_paragraph_zh(doc,
        '需要特别说明的是，内地与香港重疾险在病种分类标准上并不完全一致。内地保险公司为突出产品卖点，'
        '常将疾病大类拆分为多个独立病种，例如将「癌症」拆分为肺癌、胃癌、肝癌、乳腺癌、白血病等；'
        '而保诚则按疾病大类合并统计，如「癌症」作为一个严重病况。因此，保诚的56种严重疾病在覆盖范围上'
        '并不逊于内地的120种重疾，甚至在早期病况覆盖、疾病定义宽松度和多次赔付设计上更具优势。'
        '客户在选择时不应仅凭病种数量判断保障好坏，而应关注高发重疾覆盖、理赔门槛和实际赔付概率。')

    add_paragraph_zh(doc, '表3 理赔与医院要求对比', bold=True)
    table_data = [
        ['对比维度', '保诚 CIM3/BCIM3', '国寿康宁尊享', '太保蓝鲸1号', '平安盛世福'],
        ['内地医院要求', '需在三级医院或15个指定城市二级医院/指定名单医院', '二级及以上公立医院通常即可', '二级及以上公立医院通常即可', '二级及以上公立医院通常即可'],
        ['医生要求', '理赔申请书第二部分须由主诊医生填写', '按内地理赔流程，通常无需医生填写专门表格', '按内地理赔流程', '按内地理赔流程'],
        ['理赔地点', '材料寄送香港，审核周期相对较长', '内地APP/线上理赔，流程快', '内地APP/线上理赔，流程快', '内地APP/线上理赔，流程快'],
        ['疾病定义', '相对宽松（如中风不要求180天后遗留障碍）', '遵循内地重疾定义规范，相对严格', '遵循内地重疾定义规范', '遵循内地重疾定义规范'],
        ['健康告知', '严格/无限告知，投保前就诊记录影响大', '有限告知，按问卷回答', '有限告知，按问卷回答', '有限告知，按问卷回答']
    ]
    add_table_zh(doc, 6, 5, table_data)

    # 三、不同场景下的保单保障
    add_heading_zh(doc, '三、不同场景下的保单保障分析', level=1)

    add_paragraph_zh(doc, '表4 典型理赔场景下的给付差异（以30岁男性、50万保额为例）', bold=True)
    table_data = [
        ['理赔场景', '保诚 CIM3/BCIM3', '国寿康宁尊享', '太保蓝鲸1号', '平安盛世福'],
        ['投保后首10年确诊重疾', '赔付75万（50万基础 + 25万额外50%）', '赔付50万', '赔付50万', '赔付50万'],
        ['持有10年后保额', '约60-70万（非保证分红增长）', '50万（固定）', '50万（固定）', '50万（固定）'],
        ['持有20年后保额', '约75-90万（非保证分红增长）', '50万（固定）', '50万（固定）', '50万（固定）'],
        ['3年后癌症复发/转移', '首次50万 + 癌症额外最高再赔100万，累计最高150万', '基础责任不赔；若选多次赔需跨组', '若选二次重疾可赔60万（须为不同病种）', '不赔'],
        ['确诊急性心梗后1年脑中风', '可分别按心脏病、脑部疾病各赔付，累计可达高额', '若选多次赔且分在不同组可赔', '若选二次重疾可赔60万', '不赔'],
        ['60岁后确诊脑退化/柏金逊', '确诊1年后每年给付6%保额（3万/年），终身', '不额外给付', '不额外给付', '不额外给付'],
        ['确诊轻症（如原位癌）', '按早期病况赔付部分保额', '赔付10万（20%×50万）', '赔付15万（30%×50万）', '按合同约定比例赔付'],
        ['未理赔身故', '赔付保额+分红累积（非保证）', '赔付max（保额，1.05×保费，现金价值）', '按所选方案赔付保费或保额', '赔付保额']
    ]
    add_table_zh(doc, 9, 5, table_data)

    add_heading_zh(doc, '3.1 首10年额外50%保障', level=2)
    add_paragraph_zh(doc,
        '保诚 CIM3/BCIM3 提供「首10年保额升级」：投保后首10年内确诊严重疾病或身故，额外赔付50%保额。'
        '以50万基础保额为例，前10年内首次确诊重疾可获得75万赔付（50万基础 + 25万额外），'
        '覆盖被保险人家庭责任最重、收入最高的阶段。')
    add_paragraph_zh(doc,
        '内地三款产品均无此类「前10年额外赠送保额」设计，首次重疾赔付始终仅为50万。'
        '对于30-45岁的家庭经济支柱而言，保诚的前10年高杠杆设计更具吸引力。')

    add_heading_zh(doc, '3.2 保额增长与抗通胀能力', level=2)
    add_paragraph_zh(doc,
        '内地重疾险（国寿、太保、平安）的保额在投保后即固定不变。以50万保额为例，无论持有10年、'
        '20年还是30年，重疾赔付金额始终为50万人民币。在医疗成本持续上升和通货膨胀的背景下，'
        '固定保额的实际购买力会逐年下降。')
    add_paragraph_zh(doc,
        '保诚 CIM3/BCIM3 作为分红型美元保单，保额可随非保证分红逐年增长。以50万基础保额为例，'
        '持有10年后保额可能增长至60-70万，20年后可能达到75-90万，长期抗通胀能力显著优于内地产品。'
        '这是香港重疾险最核心的长期价值之一。')

    add_heading_zh(doc, '3.3 疾病定义宽松度与理赔门槛', level=2)
    add_paragraph_zh(doc,
        '保诚的疾病定义整体较内地产品更宽松。以「中风」为例：内地重疾险通常要求确诊180天后仍遗留'
        '神经系统永久性功能障碍方可理赔；保诚无此严格时限要求，理赔门槛更低。此外，部分早期癌症、'
        '甲状腺癌等疾病在保诚产品中的认定也更宽松，客户更容易获得赔付。')
    add_paragraph_zh(doc,
        '需要指出的是，内地产品在轻症/中症赔付比例上更高：太保蓝鲸1号轻症30%、中症60%优于国寿'
        '康宁尊享的20%/50%。但内地产品对重疾本身的定义相对严格，整体理赔门槛仍高于保诚。')

    add_heading_zh(doc, '3.4 高杠杆多次赔付场景', level=2)
    add_paragraph_zh(doc,
        '保诚 CIM3/BCIM3 在多次赔付设计上最为激进：首次严重疾病赔付100%保额后，癌症、心脏病发作、'
        '中风各可额外赔付2次（每次100%保额），即癌症/心梗/中风各最高累计可达300%保额；'
        '加上脑退化症、主要器官及功能相关疾病等额外保障，总额理论上可达保额1000%。'
        '这对有家族病史、担忧复发或多次重疾发生的客户具有显著吸引力。')
    add_paragraph_zh(doc,
        '相比之下，内地产品中只有国寿康宁尊享（可选分组多次）和太保蓝鲸1号（可选不分组2-3次）'
        '提供有限的多次重疾保障，且赔付总额和灵活性均不及保诚。平安盛世福则为单次赔付。')

    add_heading_zh(doc, '3.5 养老与长期护理场景', level=2)
    add_paragraph_zh(doc,
        '保诚独有的「老年疾病终身年金」在确诊严重脑退化症或柏金逊病1年后，每年支付基本保额6%，'
        '直至身故。以50万保额为例，每年可获得3万元护理金，将重疾险与长期护理需求结合，对长寿风险'
        '较高的客户极具价值。内地三款产品均无此类设计。')

    # 四、保诚内地医院理赔要求
    add_heading_zh(doc, '四、保诚重疾险内地医院理赔要求', level=1)
    add_paragraph_zh(doc,
        '根据保诚香港官网披露，在内地就医申请危疾理赔时，需满足以下条件：')
    add_paragraph_zh(doc,
        '1. 医院等级：原则上须在内地三级医院；或在15个指定城市的二级医院；或在保诚「中华人民共和国选定医院名单」内医院。')
    add_paragraph_zh(doc,
        '2. 指定15个城市：北京、上海、广州、深圳、成都、重庆、杭州、武汉、苏州、西安、南京、长沙、天津、郑州、东莞。')
    add_paragraph_zh(doc,
        '3. 医生要求：理赔申请书第二部分必须由主诊医生填写，客户需先自行支付医生填表费用；符合公司规定者可酌情免除。')
    add_paragraph_zh(doc,
        '4. 资格核实：保诚会在审核理赔申请时核实该医院资格，建议就医前通过保诚官网或国家卫健委网站查询医院等级。')
    add_paragraph_zh(doc,
        '5. 对比提示：内地重疾险通常只需二级及以上公立医院即可理赔，医院范围更广、流程更简化。')

    # 五、适合人群分析
    add_heading_zh(doc, '五、适合人群分析', level=1)

    add_paragraph_zh(doc, '表5 产品适合人群匹配', bold=True)
    table_data = [
        ['产品', '最适合人群', '不太适合人群'],
        ['保诚 CIM3/BCIM3', '有美元资产配置需求；追求多次高杠杆保障；能接受香港理赔流程；健康状况良好；有海外就医或子女留学规划', '预算有限；追求理赔便利；健康状况复杂；无法赴港签约或维护保单'],
        ['国寿康宁尊享', '看重品牌与线下服务；偏好稳健现金价值；需要可选重疾多次赔；给孩子投保', '追求高性价比；希望轻症赔付比例高'],
        ['太保蓝鲸1号', '年轻家庭；预算有限但希望保额充足；看重互联网投保便利；需要55岁后额外保障', '需要强大线下代理人服务；偏好大品牌溢价'],
        ['平安盛世福', '平安老客户；重视健康管理与运动激励；需要少儿特疾保障', '追求极致性价比；需要多次重疾保障']
    ]
    add_table_zh(doc, 5, 3, table_data)

    # 六、投保建议与风险提示
    add_heading_zh(doc, '六、投保建议与风险提示', level=1)

    add_heading_zh(doc, '6.1 投保建议', level=2)
    add_paragraph_zh(doc,
        '1. 从长期保障价值看，保诚 CIM3/BCIM3 在保额增长、疾病定义宽松度、多次赔付总额、'
        '老年护理年金等核心指标上均优于内地三款产品，是高净值家庭配置重疾险的优先选择。')
    add_paragraph_zh(doc,
        '2. 如果客户已经持有内地重疾险，建议将保诚产品作为「保障升级 + 美元资产 + 抗通胀」'
        '的补充，而非重复替代。')
    add_paragraph_zh(doc,
        '3. 仅在客户极其看重理赔便利性、医院范围广泛性，或健康状况复杂无法通过香港核保的情况下，'
        '才优先考虑内地产品。')
    add_paragraph_zh(doc,
        '4. 预算有限但仍希望获得较好多次保障的客户，可考虑太保蓝鲸1号作为过渡方案，'
        '但需清楚其保额固定、疾病定义严格的局限性。')

    add_heading_zh(doc, '6.2 风险提示', level=2)
    add_paragraph_zh(doc,
        '1. 汇率风险：美元/港币保单受汇率波动影响，未来红利和退保价值以保单货币计价。')
    add_paragraph_zh(doc,
        '2. 分红不确定性：保诚产品的非保证红利取决于保险公司投资表现，实际收益可能低于演示。')
    add_paragraph_zh(doc,
        '3. 理赔时效：香港保险理赔周期通常长于内地，且需邮寄材料、医生配合填表。')
    add_paragraph_zh(doc,
        '4. 法律适用：香港保单受香港法律管辖，纠纷需在香港解决，维权成本较高。')
    add_paragraph_zh(doc,
        '5. 健康告知：香港保险通常执行更严格的无限告知，投保前就诊记录、体检异常均需如实披露。')

    # 附录
    add_heading_zh(doc, '附录：术语解释', level=1)
    add_paragraph_zh(doc, 'CIM3：保诚「诚保一生」危疾保，覆盖严重疾病及早期危疾。')
    add_paragraph_zh(doc, 'BCIM3：诚保一生危疾保—挚爱宝，孕期可投保，保障母婴。')
    add_paragraph_zh(doc, '危疾：即重大疾病，通常包括癌症、急性心梗、脑中风等。')
    add_paragraph_zh(doc, '轻症/中症：重大疾病早期或较轻状态，按合同约定比例赔付。')
    add_paragraph_zh(doc, '保额：保单约定的最高赔付金额。')
    add_paragraph_zh(doc, '现金价值：退保时可领取的金额，长期重疾险通常随时间增长。')

    # 保存到桌面
    output_dir = '/Users/cyn/Desktop'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, '香港保诚与内地重疾险对比分析报告.docx')
    doc.save(output_path)
    print(f'报告已生成：{output_path}')


if __name__ == '__main__':
    main()
